from spire.doc import Document as SpireDocument, FileFormat
from docx import Document as DocxDocument
from pathlib import Path
import dtlpy as dl
import tempfile
import logging
import os

logger = logging.getLogger('pdf-to-text-logger')


class DocExtractor(dl.BaseServiceRunner):

    def doc_extraction(self, item: dl.Item, context: dl.Context) -> dl.Item:
        """
        Extracts a DOC/DOCX file item and uploads it as a TXT file.

        Args:
            context (dl.Context): Dataloop context to determine whether to extract tables from the original file.
            item (dl.Item): Dataloop item, either DOC or DOCX file.

        Returns:
            dl.Item: New text item containing the extracted content in TXT format.
        """
        node = context.node
        extract_tables = node.metadata['customNodeConfig']['extract_tables']
        remote_path_for_extractions = node.metadata['customNodeConfig']['remote_path_for_extractions']
        # Local test
        # extract_tables = False
        # remote_path_for_extractions = '/extracted_from_docs'

        suffix = Path(item.name).suffix.lower()
        if suffix not in {'.doc', '.docx'}:
            raise ValueError("Only .doc and .docx files are supported for extraction.")

        with tempfile.TemporaryDirectory() as temp_dir:
            # Download path - original items
            item_local_path = item.download(local_path=temp_dir, save_locally=True)
            logger.info(f"Downloaded item to temporary path: {item_local_path}")

            # Convert .doc to .docx if necessary
            if suffix == '.doc':
                docx_path = os.path.join(temp_dir, f"{Path(item.name).stem}.docx")
                try:
                    document = SpireDocument()
                    document.LoadFromFile(item_local_path)
                    document.SaveToFile(docx_path, FileFormat.Docx2016)
                    document.Close()

                except Exception as e:
                    raise RuntimeError(f"Error converting item {item.id} to .docx format: {e}")

            else:
                docx_path = item_local_path

            text = self.extract_content(docx_path=docx_path, extract_tables=extract_tables)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = os.path.join(temp_dir, f"{Path(item.name).stem}_text.txt")
            with open(output_path, "w", encoding="utf-8") as temp_text_file:
                temp_text_file.write(text)
            logger.info(f"Text saved to temporary file: {output_path}")

            new_item = item.dataset.items.upload(local_path=output_path,
                                                 remote_path=remote_path_for_extractions,
                                                 item_metadata={
                                                     'user': {'extracted_from_docs': True,
                                                              'original_item_id': item.id}})

            if new_item is None:
                raise dl.PlatformException(f"No items was uploaded! local paths: {output_path}")

        return new_item

    @staticmethod
    def extract_content(docx_path, extract_tables=True) -> str:
        """
        Extracts text content from a DOCX file, including optional table content,
        and saves the extracted text to a txt file.

        Args:
            docx_path (str): The local path to the DOCX file to be processed.
            extract_tables (bool, optional): Whether to extract text from tables in the DOCX file. Default is True.

        Returns:
            str: The text extracted.
        """
        doc = DocxDocument(docx_path)
        full_text = list()
        table_index = 0
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Check if element is a paragraph
                para = element.text
                full_text.append(para)

            elif element.tag.endswith('tbl') and extract_tables is True:  # Check if element is a table
                table = doc.tables[table_index]
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    full_text.append(row_text)
                table_index += 1

                # Join all collected text into one string
        text = '\n'.join(full_text)

        return text
